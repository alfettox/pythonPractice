import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
from tkinter import font
from docx import Document
from docx2pdf import convert

class DocxEditor:
    def __init__(self, root):
        self.root = root
        self.document = None
        self.selected_line = None
        self.selected_font = None

        self.text_box = tk.Text(self.root, height=20, width=50)
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.scrollbar = ttk.Scrollbar(self.root, orient=tk.VERTICAL, command=self.text_box.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_box.config(yscrollcommand=self.scrollbar.set)

        self.open_button = tk.Button(self.root, text="Open", command=self.open_document)
        self.open_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.save_button = tk.Button(self.root, text="Save", command=self.save_document)
        self.save_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.align_left_button = tk.Button(self.root, text="Align Left", command=self.align_left)
        self.align_left_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.align_right_button = tk.Button(self.root, text="Align Right", command=self.align_right)
        self.align_right_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.font_var = tk.StringVar()
        self.font_dropdown = ttk.Combobox(self.root, textvariable=self.font_var, state="readonly")
        self.font_dropdown.pack(side=tk.LEFT, padx=5, pady=5)
        self.font_dropdown.bind("<<ComboboxSelected>>", self.change_font)
        self.populate_font_dropdown()

        self.save_pdf_button = tk.Button(self.root, text="Save as PDF", command=self.save_as_pdf)
        self.save_pdf_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.text_box.bind("<Button-1>", self.select_line)

    def open_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if file_path:
            self.document = Document(file_path)
            self.display_document()

    def save_document(self):
        if self.document:
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
            if file_path:
                self.document.save(file_path)

    def display_document(self):
        if self.document:
            content = []
            for paragraph in self.document.paragraphs:
                content.append(paragraph.text)
            self.text_box.delete("1.0", tk.END)
            self.text_box.insert(tk.END, "\n".join(content))
            self.update_text_font()

    def align_left(self):
        if self.document and self.selected_line:
            self.selected_line.alignment = 0  # 0 for left alignment
            self.display_document()

    def align_right(self):
        if self.document and self.selected_line:
            self.selected_line.alignment = 2  # 2 for right alignment
            self.display_document()

    def select_line(self, event):
        self.text_box.tag_remove(tk.SEL, "1.0", tk.END)
        line_index = self.text_box.index(tk.CURRENT).split(".")[0]
        if line_index and self.document:
            self.selected_line = self.document.paragraphs[int(line_index)]
            self.text_box.tag_add(tk.SEL, f"{line_index}.0", f"{line_index}.end")

    def populate_font_dropdown(self):
        available_fonts = font.families()
        self.font_dropdown["values"] = available_fonts
        if available_fonts:
            self.font_dropdown.current(0)

    def change_font(self, event):
        selected_font = self.font_var.get()
        if selected_font:
            self.selected_font = selected_font
            self.update_text_font()

    def update_text_font(self):
        if self.selected_font:
            current_font = font.Font(family=self.selected_font, size=12)
            self.text_box.configure(font=current_font)

    def save_as_pdf(self):
        if self.document:
            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
            if file_path:
                convert(file_path, self.document)

root = tk.Tk()
root.title("DOCX Editor")
editor = DocxEditor(root)
root.mainloop()
